import pytest
import shutil
import os
from docx import Document
from django.contrib.auth.models import User
from itempool.models import ItemPool, ImportBatch, DraftItem
from itempool.services.import_docx import DocxImportService

@pytest.mark.django_db
def test_docx_import_unlabeled_questions(tmp_path):
    # Kullanıcı ve havuz oluştur
    user = User.objects.create_user(username='testuser', password='password')
    pool = ItemPool.objects.create(name='Test Pool', owner=user)
    
    # Word belgesi oluştur
    doc = Document()
    doc.add_paragraph('Aşağıdakilerden hangisi tipik performans testidir?')
    doc.add_paragraph('A) Genel zekayı ölçen bir IQ testi')
    doc.add_paragraph('B) Matematiksel becerileri değerlendiren bir test')
    doc.add_paragraph('C) Sınav kaygısını ölçen bir test')
    doc.add_paragraph('D) Atletler için fiziksel uygunluk testi')
    doc.add_paragraph('E) Ders başarısını değerlendiren bir final sınavı')
    doc.add_paragraph('Cevap: D')
    
    doc_path = tmp_path / "test_questions.docx"
    doc.save(doc_path)
    
    # Batch oluştur
    from django.core.files.uploadedfile import SimpleUploadedFile
    with open(doc_path, "rb") as f:
        docx_file = SimpleUploadedFile("test_questions.docx", f.read(), content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
    batch = ImportBatch.objects.create(
        pool=pool,
        original_filename='test_questions.docx',
        uploaded_file=docx_file,
        created_by=user
    )
    
    # Parser çalıştır
    service = DocxImportService(batch.id)
    count = service.process()
    
    # Doğrula: 1 soru üretilmiş olmalı (seçenekler ayrı soru olmamalı!)
    assert count == 1
    
    drafts = DraftItem.objects.filter(batch=batch)
    assert drafts.count() == 1
    
    draft = drafts.first()
    assert draft.correct_answer == 'D'
    assert len(draft.choices_json) == 5
    assert draft.choices_json[0]['label'] == 'A'
    assert "IQ testi" in draft.choices_json[0]['text']
    assert draft.choices_json[3]['label'] == 'D'
    assert "fiziksel uygunluk" in draft.choices_json[3]['text']
