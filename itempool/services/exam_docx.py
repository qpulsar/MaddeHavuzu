"""
Sınav kağıdı Word (.docx) üretim servisi.
TestForm + ExamTemplate → python-docx → .docx bytes
"""
import io
from datetime import date
from docx import Document
from docx.shared import Pt, Mm, Twips
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _set_columns(section, column_count, show_divider=False):
    """Bölümün sütun sayısını ayarlar."""
    section.start_type = WD_SECTION.CONTINUOUS
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(column_count))
    if show_divider:
        cols.set(qn('w:sep'), '1')
    cols.set(qn('w:space'), '708')  # 0.5 inch approx

def _resolve_variable(text: str, context: dict) -> str:
    """Şablon metin içindeki {variable} alanlarını doldurur."""
    if not text: return ""
    for key, val in context.items():
        text = text.replace(f'{{{key}}}', str(val))
    return text

def generate_exam_docx(test_form, template: "ExamTemplate", with_answer_key: bool = False) -> io.BytesIO:
    """
    Sınav formunu Word (.docx) formatında üretir.
    """
    doc = Document()
    
    # 1. Sayfa Ayarları
    section = doc.sections[0]
    # Kağıt boyutu (Basit eşleştirme)
    if template.page_size == 'A4':
        section.page_height = Mm(297)
        section.page_width = Mm(210)
    elif template.page_size == 'A5':
        section.page_height = Mm(210)
        section.page_width = Mm(148)
    
    section.top_margin = Mm(template.margin_top)
    section.bottom_margin = Mm(template.margin_bottom)
    section.left_margin = Mm(template.margin_left)
    section.right_margin = Mm(template.margin_right)

    # 2. Üst Bilgi (Header) - Word'de 3 sütunlu tablo en iyisidir
    var_context = {
        'form_name': test_form.name,
        'course': test_form.course.name if test_form.course else 'Genel',
        'course_code': test_form.course.code if test_form.course else '—',
        'semester': test_form.course.semester if test_form.course else 'Genel',
        'teacher_name': test_form.created_by.get_full_name() if test_form.created_by else '—',
        'date': date.today().strftime('%d.%m.%Y'),
        'page': '1', # Word dynamic field insertion is complex, using static placeholder
        'total_pages': '?',
    }

    # Font ayarları (Default)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(template.font_size)

    # Başlık — GrapesJS HTML'den metin ayıklama
    import re
    def clean_html(html):
        if not html: return ""
        html = re.sub(r'<(style|script)[^>]*>.*?</\1>', '', html, flags=re.DOTALL)
        html = re.sub(r'<(p|br|div|tr|h[1-6])[^>]*>', '\n', html)
        text = re.sub(r'<[^>]+>', '', html)
        text = text.replace('&nbsp;', ' ').replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
        return "\n".join([line.strip() for line in text.split('\n') if line.strip()])

    if template.header_html:
        resolved_html = _resolve_variable(template.header_html, var_context)
        header_text = clean_html(resolved_html)
        for line in header_text.split('\n'):
            p_h = doc.add_paragraph(line)
            p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_h.paragraph_format.space_after = Pt(0)
    else:
        # Boş başlık — sınav adını yaz
        p_header = doc.add_paragraph(var_context.get('form_name', ''))
        p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_header.runs[0].bold = True if p_header.runs else None

    if template.show_header_line:
        doc.add_paragraph().add_run("_" * 80).bold = True


    form_items = test_form.form_items.select_related(
        'item_instance__item',
        'item_instance__learning_outcome'
    ).prefetch_related('item_instance__item__choices').order_by('order')

    # 3. Öğrenci Bilgi Kutusu
    if not with_answer_key and template.show_student_info_box:
        box_table = doc.add_table(rows=1, cols=3)
        box_table.style = 'Table Grid'
        b_cells = box_table.rows[0].cells
        b_cells[0].text = "Ad Soyad:"
        b_cells[1].text = "No:"
        b_cells[2].text = "İmza:"
        doc.add_paragraph() # Spacer

    # 4. Sütun Ayarı (Sorular için)
    if not with_answer_key and template.column_count > 1:
        _set_columns(section, template.column_count, template.column_divider)

    # 5. Sorular
    if not with_answer_key:
        for fi in form_items:
            item = fi.item_instance.item
            
            # Soru Kökü
            p_stem = doc.add_paragraph()
            run_num = p_stem.add_run(f"{fi.order}. ")
            run_num.bold = True
            
            p_stem.add_run(item.stem)
            
            if template.show_question_points:
                run_pts = p_stem.add_run(f" ({fi.points} puan)")
                run_pts.italic = True
                run_pts.font.size = Pt(8)
                
            # Kazanımlar
            outcome = fi.item_instance.learning_outcome
            if outcome:
                run_outcomes = p_stem.add_run(f" ({outcome.code})")
                run_outcomes.italic = True
                run_outcomes.font.size = Pt(8)
            
            # Şıklar (MCQ/TF)
            if item.item_type in ['MCQ', 'TF']:
                if fi.choice_overrides:
                    choices = fi.choice_overrides
                else:
                    choices = [{'label': c.label, 'text': c.text} for c in item.choices.all()]
                
                # Layout seçimi (Sayfa sütununa göre daraltılmış eşikler)
                max_len = max([len(str(c.get('text', ''))) for c in choices]) if choices else 0
                col_factor = template.column_count
                
                t_vert = 45 if col_factor == 1 else (28 if col_factor == 2 else 18)
                t_grid3 = 20 if col_factor == 1 else (10 if col_factor == 2 else 6)
                
                # 1 Sütun (Vertical)
                if max_len > t_vert or template.choice_layout == 'vertical':
                    for c in choices:
                        p_choice = doc.add_paragraph()
                        p_choice.paragraph_format.left_indent = Pt(15)
                        p_choice.paragraph_format.space_after = Pt(template.choice_spacing)
                        p_choice.add_run(f"{c['label']}) ").bold = True
                        p_choice.add_run(str(c['text']))
                
                # Grid (2 veya 3 Sütun)
                else:
                    cols_count = 3 if max_len < t_grid3 else 2
                    choice_table = doc.add_table(rows=0, cols=cols_count)
                    for i in range(0, len(choices), cols_count):
                        row_cells = choice_table.add_row().cells
                        for j in range(cols_count):
                            if i + j < len(choices):
                                c = choices[i+j]
                                row_cells[j].text = f"{c['label']}) {c['text']}"
                    
                    # Spacer
                    p_spacer = doc.add_paragraph()
                    p_spacer.paragraph_format.space_before = Pt(0)
                    p_spacer.paragraph_format.space_after = Pt(template.question_spacing)
            
            elif item.item_type == 'SHORT_ANSWER':
                p_sa = doc.add_paragraph("_________________________________________________")
                p_sa.paragraph_format.space_after = Pt(template.question_spacing)
            else:
                # Diğer tipler veya boşluk
                p_space = doc.add_paragraph()
                p_space.paragraph_format.space_before = Pt(0)
                p_space.paragraph_format.space_after = Pt(template.question_spacing)

    # 6. Cevap Anahtarı
    if with_answer_key:
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(f"CEVAP ANAHTARI — {test_form.name}")
        run_title.bold = True
        run_title.font.size = Pt(14)
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_title.paragraph_format.space_after = Pt(12)
        
        # 5 sütunlu grid tablo oluşturalım
        cols_count = 5
        key_table = doc.add_table(rows=0, cols=cols_count)
        key_table.style = 'Table Grid'
        
        import math
        num_items = len(form_items)
        num_rows = math.ceil(num_items / cols_count)
        
        grid_data = []
        for fi in form_items:
            # Doğru şıkkı bul
            correct_ans = ""
            choices = fi.get_choices()
            if choices:
                correct_choice = next((c for c in choices if c.get('is_correct')), None)
                if correct_choice:
                    correct_ans = correct_choice.get('label', '')
            
            if not correct_ans:
                item = fi.item_instance.item
                if item.item_type == 'SHORT_ANSWER':
                    correct_ans = item.expected_answer or "Kısa Cevap"
                elif item.item_type == 'OPEN':
                    correct_ans = "Açık"
                else:
                    correct_ans = "-"
            
            grid_data.append(f"{fi.order}. {correct_ans}")
            
        for r in range(num_rows):
            row_cells = key_table.add_row().cells
            for c in range(cols_count):
                idx = r * cols_count + c
                if idx < len(grid_data):
                    row_cells[c].text = grid_data[idx]
                    row_cells[c].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    row_cells[c].text = ""

    # Çıktı
    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream
